"""
download_fulltext.py —— 导出"用于分析的案件"判决原文（通用 / 证券 两套工作流共用）
─────────────────────────────────────────────
用法：python3 scripts/download_fulltext.py <research_dir> [--docx]
输入：
  <research_dir>/05_enriched_cases.json   分析池（通用=全部普通案例；证券=已折叠的核心判决/示范判决）
  <research_dir>/03_raw_cases.json        字段回退（按 Gid）
输出：
  <research_dir>/output/原文/<序号>_<案号或标题>.md   每案一份判决原文
  <research_dir>/output/原文/00_索引.md               清单（案号/法院/等级/是否有正文/链接）
  （加 --docx 时，额外用 build_report_docx 的依赖把每份 .md 渲染成 .docx；无 python-docx 则自动跳过）

原理：MCP 的 Ascertain/Identified/RefereeBasis/RefereeResult/PlaintiffClaims/DefenseViewpoint/TrialAfter
      就是法宝对判决书正文的结构化拆解。本脚本按判决书结构把这些字段拼回一份原文，全部来自 MCP，
      不抓取 pkulaw 网页（付费/登录），每份附 pkulaw 链接作权威溯源。
说明：只有普通案例(CaseGrade=07)的正文字段非空；典型/参考/经典案例正文为空，仅导出元数据 + 链接。
依赖：标准库即可；--docx 需 pip3 install python-docx。
"""

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent / "common"))
from pkulaw_utils import (  # noqa: E402
    load, clean_url, flatten_court, flatten_leaf, field, build_raw_index,
)

# 判决书正文字段 → 章节标题（按判决书自然顺序）
SECTIONS = [
    ("PlaintiffClaims", "一、原告诉讼请求"),
    ("DefenseViewpoint", "二、被告答辩与抗辩"),
    ("TrialAfter", "三、审理经过"),
    ("Ascertain", "四、经审理查明的事实"),
    ("ControversialFocus", "五、争议焦点"),
    ("Identified", "六、本院认为"),
    ("RefereeBasis", "七、裁判依据"),
    ("RefereeResult", "八、裁判结果"),
]


def safe_name(s: str, fallback: str) -> str:
    """文件名安全化：去掉路径分隔符与多余空白，限长。"""
    s = (s or "").strip() or fallback
    s = re.sub(r"[\\/:*?\"<>|\n\r\t]+", "_", s)
    return s[:60].rstrip("_ ") or fallback


def has_body(c, raw_idx) -> bool:
    return any(field(c, raw_idx, k) for k, _ in SECTIONS if k in
              ("Ascertain", "Identified", "RefereeBasis", "RefereeResult"))


def build_one_md(i, c, raw_idx) -> tuple:
    """返回 (文件名, markdown文本, 是否有正文)。"""
    title = field(c, raw_idx, "Title")
    case_flag = field(c, raw_idx, "CaseFlag")
    _, court = flatten_court(field(c, raw_idx, "LastInstanceCourt", {}))
    grade = flatten_leaf(field(c, raw_idx, "CaseGrade", {})) or "普通案例"
    url = clean_url(field(c, raw_idx, "Url"))
    body = has_body(c, raw_idx)

    lines = [f"# {title}", ""]
    meta = [
        ("案号", case_flag),
        ("审理法院", court),
        ("案由", flatten_leaf(field(c, raw_idx, "Category", {}))),
        ("审级", field(c, raw_idx, "CaseClassName")),
        ("裁判日期", field(c, raw_idx, "LastInstanceDate")),
        ("案件等级", grade),
    ]
    # 证券工作流的附加标识（存在才写）
    for k in ("涉案上市公司", "虚假陈述事件", "核心判决类型"):
        v = c.get(k)
        if v:
            meta.append((k, v))
    for key, val in meta:
        if val:
            lines.append(f"- **{key}**：{val}")
    if url:
        lines.append(f"- **北大法宝原文链接**：{url}")
    lines.append("")
    lines.append("---")
    lines.append("")

    if body:
        for key, heading in SECTIONS:
            text = field(c, raw_idx, key)
            if not text:
                continue
            lines.append(f"## {heading}")
            lines.append("")
            lines.append(str(text).strip())
            lines.append("")
    else:
        lines.append("> ⚠ 该案为典型/参考/经典案例，北大法宝未返回判决书正文要素（正文字段为空），"
                     "此处仅有元数据。完整原文请点击上方北大法宝链接查看。")
        lines.append("")

    fname = f"{i:03d}_{safe_name(case_flag or title, f'case{i}')}.md"
    return fname, "\n".join(lines), body


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    want_docx = "--docx" in sys.argv
    if not args:
        sys.exit("用法：python3 scripts/download_fulltext.py <research_dir> [--docx]")
    rd = Path(args[0])
    cases = load(rd / "05_enriched_cases.json")
    if not cases:
        sys.exit(f"未找到或为空：{rd/'05_enriched_cases.json'}")
    raw_idx = build_raw_index(rd)

    out_dir = rd / "output" / "原文"
    out_dir.mkdir(parents=True, exist_ok=True)

    index_rows, n_body, n_meta = [], 0, 0
    md_paths = []
    for i, c in enumerate(cases, 1):
        fname, md, body = build_one_md(i, c, raw_idx)
        (out_dir / fname).write_text(md, encoding="utf-8")
        md_paths.append(out_dir / fname)
        n_body += 1 if body else 0
        n_meta += 0 if body else 1
        index_rows.append((i, fname, field(c, raw_idx, "CaseFlag"),
                           flatten_court(field(c, raw_idx, "LastInstanceCourt", {}))[1],
                           flatten_leaf(field(c, raw_idx, "CaseGrade", {})) or "普通案例",
                           "有正文" if body else "仅元数据",
                           clean_url(field(c, raw_idx, "Url"))))

    idx = ["# 分析案件原文索引", "",
           f"共 {len(cases)} 份（有正文 {n_body} 份 · 仅元数据 {n_meta} 份）。原文要素均来自北大法宝 MCP。", "",
           "| 序号 | 文件 | 案号 | 审理法院 | 案件等级 | 正文 | 链接 |",
           "|---|---|---|---|---|---|---|"]
    for r in index_rows:
        idx.append(f"| {r[0]} | {r[1]} | {r[2]} | {r[3]} | {r[4]} | {r[5]} | {r[6]} |")
    (out_dir / "00_索引.md").write_text("\n".join(idx), encoding="utf-8")

    print(f"原文已导出：{out_dir}")
    print(f"  共 {len(cases)} 份 | 有正文 {n_body} 份 | 仅元数据 {n_meta} 份 | 索引 00_索引.md")

    if want_docx:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            print("  （--docx 跳过：未安装 python-docx，pip3 install python-docx 后重试）")
            return
        for p in md_paths:
            doc = Document()
            for line in p.read_text(encoding="utf-8").splitlines():
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=0)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=1)
                elif line.strip() in ("---", ""):
                    continue
                else:
                    doc.add_paragraph(line)
            doc.save(str(p.with_suffix(".docx")))
        print(f"  已额外生成 {len(md_paths)} 份 .docx")


if __name__ == "__main__":
    main()
